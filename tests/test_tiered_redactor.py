# tests for the tiered detection pipeline

import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document

from pii_redactor.docx_redaction_service import DocxRedactionService
from pii_redactor.models import EntitySpan, PipelineMode
from pii_redactor.scanners import (
    EntropyDetector,
    ExclusionFilter,
    NerDetector,
    RegexDetector,
    TieredPipelineScanner,
    calculate_shannon_entropy,
)


def local_pipeline():
    return TieredPipelineScanner(scanners=[RegexDetector(), EntropyDetector(), NerDetector(use_optional_ner=False)])


# --- tier 1: regex based stuff ---
class Tier1RegexDetectorTests(unittest.TestCase):
    def setUp(self):
        self.detector = RegexDetector()

    # throw a bunch of PII at it and see what sticks
    def test_structured_entities(self):
        text = (
            "Contact info: email user@example.com, URL https://example.com/api, "
            "IP 192.168.1.50, SSN 123-45-6789, Card 4111 1111 1111 1111"
        )
        spans = self.detector.scan(text)
        types = {s.entity_type: s.text for s in spans}

        self.assertIn("email", types)
        self.assertEqual(types["email"], "user@example.com")
        self.assertIn("url", types)
        self.assertEqual(types["url"], "https://example.com/api")
        self.assertIn("ip_address", types)
        self.assertEqual(types["ip_address"], "192.168.1.50")
        self.assertIn("ssn", types)
        self.assertEqual(types["ssn"], "123-45-6789")
        self.assertIn("credit_card", types)
        self.assertEqual(types["credit_card"], "4111 1111 1111 1111")


# --- tier 2: entropy / secret detection ---
class Tier2EntropyDetectorTests(unittest.TestCase):
    def setUp(self):
        self.detector = EntropyDetector()

    # basic sanity check - random looking strings should have higher entropy
    def test_shannon_entropy_calculation(self):
        low_entropy = calculate_shannon_entropy("aaaaaaaaaaaa")
        high_entropy = calculate_shannon_entropy("fk_test_94f8a12b0e9c3d4f")
        self.assertGreater(high_entropy, low_entropy)

    # TODO: add more key formats (AWS, GCP, etc)
    def test_token_and_api_key_detection(self):
        text = "API_KEY=fk_test_94f8a12b0e9c3d4f8a12b0e9 and JWT eyJhbGciOiJIUzI1NiJ9.eyJzdWIiOiIxMjM0NTY3ODkwIn0.SflKxwRJSMeKKF2QT4fwpMeJf36POk6yJV_adQssw5c"
        spans = self.detector.scan(text)
        detected_types = [s.entity_type for s in spans]

        self.assertTrue("api_key" in detected_types or "token" in detected_types)


# --- tier 3: context-based NER ---
class Tier3NerDetectorTests(unittest.TestCase):
    def setUp(self):
        self.detector = NerDetector(use_optional_ner=False)

    # names after keywords like "Contact Person:" should get detected
    def test_context_ner_person_detection(self):
        text = "OUR PROMOTERS: AARAV MEHTA. Contact Person: Rajesh Kumar."
        spans = self.detector.scan(text)
        persons = [s.text for s in spans if s.entity_type == "person"]

        self.assertTrue(any("AARAV MEHTA" in p for p in persons))
        self.assertTrue(any("Rajesh Kumar" in p for p in persons))


# public orgs like BSE/SEBI shouldn't be redacted
class ExclusionFilterTests(unittest.TestCase):
    def setUp(self):
        self.filter = ExclusionFilter(public_organizations={"BSE", "SEBI", "Reserve Bank of India"})

    def test_excludes_public_organizations(self):
        public_span = EntitySpan("company", "BSE", 0, 3, 0.9, "test")
        private_span = EntitySpan("company", "Nuvama Wealth Management Limited", 0, 32, 0.9, "test")

        self.assertTrue(self.filter.is_excluded(public_span))
        self.assertFalse(self.filter.is_excluded(private_span))

        filtered = self.filter.filter_spans([public_span, private_span])
        self.assertEqual(len(filtered), 1)
        self.assertEqual(filtered[0].text, "Nuvama Wealth Management Limited")


# end-to-end check that all tiers work together
class TieredPipelineScannerTests(unittest.TestCase):
    def setUp(self):
        self.pipeline = local_pipeline()

    def test_end_to_end_tiered_scan(self):
        text = (
            "Email: test@domain.com, API Key: api_key=fk_test_94f8a12b0e9c3d4f8a12b0e9, "
            "Contact Person: Lokesh Shah."
        )
        spans = self.pipeline.scan_text(text)
        types = {s.entity_type for s in spans}

        self.assertIn("email", types)
        self.assertTrue("person" in types or "company" in types or "api_key" in types)


class PipelineModeTests(unittest.TestCase):
    # mutation mode should actually change the docx
    def test_mutation_mode_rewrites_document(self):
        with TemporaryDirectory() as tmp:
            docx_path = Path(tmp) / "input.docx"
            out_path = Path(tmp) / "out_mutation.docx"

            doc = Document()
            doc.add_paragraph("Contact email is user@example.com.")
            doc.save(docx_path)

            service = DocxRedactionService(scanner=local_pipeline(), mode=PipelineMode.MUTATION)
            audit = service.redact_docx(docx_path, out_path)

            self.assertGreater(len(audit), 0)
            self.assertTrue(out_path.exists())

            redacted_doc = Document(out_path)
            self.assertNotIn("user@example.com", redacted_doc.paragraphs[0].text)

    # validation mode should leave the original file alone
    def test_validation_mode_does_not_mutate_document(self):
        with TemporaryDirectory() as tmp:
            docx_path = Path(tmp) / "input.docx"

            doc = Document()
            doc.add_paragraph("Contact email is user@example.com.")
            doc.save(docx_path)

            service = DocxRedactionService(scanner=local_pipeline(), mode=PipelineMode.VALIDATION)
            audit = service.redact_docx(docx_path, None)

            self.assertGreater(len(audit), 0)

            original_doc = Document(docx_path)
            self.assertIn("user@example.com", original_doc.paragraphs[0].text)


if __name__ == "__main__":
    unittest.main()
